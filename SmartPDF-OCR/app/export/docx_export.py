"""
DOCX 导出模块
将 OCR 结果导出为 Word 文档
"""

from pathlib import Path
from typing import List, Union, Optional
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.ocr.postprocess import ProcessedPage


class DocxExporter:
    """DOCX 导出器"""
    
    def __init__(
        self,
        font_name: str = "宋体",
        font_size: int = 12,
        line_spacing: float = 1.5,
        page_width: float = 21.0,  # A4 宽度 cm
        page_height: float = 29.7,  # A4 高度 cm
        margin: float = 2.54  # 页边距 cm
    ):
        """
        初始化 DOCX 导出器
        
        Args:
            font_name: 字体名称
            font_size: 字体大小（磅）
            line_spacing: 行间距
            page_width: 页面宽度（cm）
            page_height: 页面高度（cm）
            margin: 页边距（cm）
        """
        self.font_name = font_name
        self.font_size = font_size
        self.line_spacing = line_spacing
        self.page_width = page_width
        self.page_height = page_height
        self.margin = margin
    
    def export(
        self,
        pages: List[ProcessedPage],
        output_path: Union[str, Path],
        title: Optional[str] = None,
        include_page_breaks: bool = True,
        is_markdown: bool = False
    ) -> Path:
        """
        导出为 DOCX 文件
        
        Args:
            pages: 处理后的页面列表
            output_path: 输出文件路径
            title: 文档标题
            include_page_breaks: 是否在每页后添加分页符
            is_markdown: 是否按 Markdown 格式渲染内容
            
        Returns:
            输出文件路径
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        doc = Document()
        
        # 设置页面尺寸
        self._setup_page(doc)
        
        # 添加标题
        if title:
            self._add_title(doc, title)
        
        # 添加内容
        for i, page in enumerate(pages):
            if is_markdown:
                self._add_markdown_content(doc, page)
            else:
                self._add_page_content(doc, page)
            
            # 添加分页符
            if include_page_breaks and i < len(pages) - 1:
                doc.add_page_break()
        
        doc.save(str(output_path))
        return output_path
    
    def _setup_page(self, doc: Document) -> None:
        """设置页面"""
        section = doc.sections[0]
        section.page_width = Cm(self.page_width)
        section.page_height = Cm(self.page_height)
        section.left_margin = Cm(self.margin)
        section.right_margin = Cm(self.margin)
        section.top_margin = Cm(self.margin)
        section.bottom_margin = Cm(self.margin)
    
    def _add_title(self, doc: Document, title: str) -> None:
        """添加标题"""
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_page_content(self, doc: Document, page: ProcessedPage) -> None:
        """添加普通页面内容"""
        for paragraph in page.paragraphs:
            p = doc.add_paragraph()
            self._add_run_with_style(p, paragraph.text)
            p.paragraph_format.line_spacing = self.line_spacing

    def _add_markdown_content(self, doc: Document, page: ProcessedPage) -> None:
        """解析并添加 Markdown 内容"""
        import re
        
        # 合并所有段落文本（通常 AI 结果是一大段）
        full_text = "\n".join([p.text for p in page.paragraphs])
        lines = full_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 标题 (Headers)
            if line.startswith('#'):
                # 计算级别
                level = 0
                for char in line:
                    if char == '#':
                        level += 1
                    else:
                        break
                content = line[level:].strip()
                if level > 9: level = 9
                doc.add_heading(content, level=level)
                continue
            
            # 列表 (Lists)
            if line.startswith('- ') or line.startswith('* '):
                content = line[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
            elif re.match(r'^\d+\.\s', line):
                # 简单有序列表检测
                match = re.match(r'^(\d+\.\s)(.*)', line)
                content = match.group(2)
                p = doc.add_paragraph(style='List Number')
            else:
                p = doc.add_paragraph()
                content = line
                p.paragraph_format.line_spacing = self.line_spacing
            
            # 内联格式处理 (Bold)
            # 分割 **bold**
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**') and len(part) > 4:
                    self._add_run_with_style(p, part[2:-2], bold=True)
                else:
                    self._add_run_with_style(p, part)

    def _add_run_with_style(self, paragraph, text, bold=False):
        """添加带样式的文本块"""
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        run._element.rPr.rFonts.set(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia',
            self.font_name
        )
    
    def export_with_confidence(
        self,
        pages: List[ProcessedPage],
        output_path: Union[str, Path],
        highlight_threshold: float = 0.5
    ) -> Path:
        """
        导出并标记低置信度内容
        
        Args:
            pages: 处理后的页面列表
            output_path: 输出文件路径
            highlight_threshold: 高亮阈值
            
        Returns:
            输出文件路径
        """
        from docx.shared import RGBColor
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        doc = Document()
        self._setup_page(doc)
        
        for page in pages:
            for paragraph in page.paragraphs:
                p = doc.add_paragraph()
                
                # 检查置信度
                if paragraph.avg_confidence < highlight_threshold:
                    # 低置信度用红色标记
                    run = p.add_run(paragraph.text)
                    run.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    run = p.add_run(paragraph.text)
                
                run.font.name = self.font_name
                run.font.size = Pt(self.font_size)
        
        doc.save(str(output_path))
        return output_path


def export_to_docx(
    pages: List[ProcessedPage],
    output_path: Union[str, Path],
    title: Optional[str] = None
) -> Path:
    """
    快捷函数：导出为 DOCX 文件
    
    Args:
        pages: 处理后的页面列表
        output_path: 输出文件路径
        title: 文档标题
        
    Returns:
        输出文件路径
    """
    exporter = DocxExporter()
    return exporter.export(pages, output_path, title=title)
